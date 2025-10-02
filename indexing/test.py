from fastapi import FastAPI, HTTPException, Query, UploadFile, File
from pydantic import BaseModel
from typing import List, Optional
from elasticsearch import Elasticsearch
from sentence_transformers import SentenceTransformer
import os
from dotenv import load_dotenv
from docx import Document as DocxDocument
import uvicorn

load_dotenv()
ELASTICSEARCH_ENDPOINT = os.getenv("ELASTICSEARCH_ENDPOINT")
PORT = int(os.getenv("API_SEVER_PORT"))

# --- ElasticSearch 接続 ---
es = Elasticsearch(ELASTICSEARCH_ENDPOINT)

# --- Hugging Face 埋め込みモデル ---
embedding_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
VECTOR_DIM = embedding_model.get_sentence_embedding_dimension()

# --- FastAPI 初期化 ---
app = FastAPI(title="RAG Document Indexing API")

# --- Pydantic モデル ---
class IndexRequest(BaseModel):
    index_name: str
    description: Optional[str] = None

class Document(BaseModel):
    description: str
    content: str

class ChunkedDocument(BaseModel):
    index_name: str
    description: str
    content: str
    chunk_size: Optional[int] = 200  # 1チャンクの文字数デフォルト200

class DocumentChunkRequest(BaseModel):
    index_name: str                # 保存先インデックス名
    description: str               # ドキュメントの説明 (検索対象フィールド)
    content: str                   # 登録する元テキスト
    chunk_size: int = 200          # チャンクサイズ（デフォルト200文字）

# --- ファイル前処理関数 ---
def preprocess_text(text: str, chunk_size: int = 500) -> List[dict]:
    text = text.replace("\r\n", "\n").strip()
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    indexed_docs = []
    for i, chunk in enumerate(chunks):
        vector = embedding_model.encode(chunk).tolist()
        indexed_docs.append({"content": chunk, "embedding": vector, "chunk_number": i+1})
    return indexed_docs

def preprocess_file(file_path: str, chunk_size: int = 500) -> List[dict]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    elif ext == ".docx":
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
    else:
        raise ValueError("対応していないファイル形式です。txtかdocxを使用してください。")
    return preprocess_text(text, chunk_size)

# --- インデックス作成 ---
# @app.post("/create_index/")
# def create_index(request: IndexRequest):
#     index_body = {
#         "settings": {
#             "index": {
#                 "number_of_shards": 1,
#                 "number_of_replicas": 0
#             }
#         },
#         "mappings": {
#             "properties": {
#                 "description": {"type": "text"},
#                 "content": {"type": "text"},
#                 "embedding": {"type": "dense_vector", "dims": VECTOR_DIM}
#             }
#         }
#     }

#     try:
#         es.indices.create(index=request.index_name, body=index_body, ignore=400)
#         return {
#             "message": f"Index '{request.index_name}' created successfully.",
#             "index_description": request.description or "No description"
#         }
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

@app.post("/create_index/")
def create_index(request: IndexRequest):
    index_body = {
        "settings": {
            "index": {
                "number_of_shards": 1,
                "number_of_replicas": 0
            }
        },
        "mappings": {
            "properties": {
                "description": {"type": "text"},
                "content": {"type": "text"},
                "embedding": {"type": "dense_vector", "dims": VECTOR_DIM}
            }
        }
    }

    try:
        es.indices.create(index=request.index_name, body=index_body, ignore=400)
        
        # --- _meta_ ドキュメント登録 ---
        es.index(
            index=request.index_name,
            id="_meta_",
            document={"description": request.description or ""}
        )
        
        return {
            "message": f"Index '{request.index_name}' created successfully.",
            "index_description": request.description or "No description"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- ドキュメント登録（チャンク化対応） ---
@app.post("/index_document_chunked/")
def index_document_chunked(request: DocumentChunkRequest):
    try:
        chunks = [request.content[i:i+request.chunk_size] for i in range(0, len(request.content), request.chunk_size)]
        for chunk in chunks:
            vector = embedding_model.encode(chunk).tolist()
            doc = {
                "description": request.description,
                "content": chunk,
                "embedding": vector
            }
            es.index(index=request.index_name, body=doc)
        return {"message": f"Document indexed into '{request.index_name}' in {len(chunks)} chunks"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- ファイルアップロードからインデックスに登録するエンドポイント ---
@app.post("/index_file/")
async def index_file(file: UploadFile = File(...), index_name: str = "rag_docs", chunk_size: int = 500):
    tmp_path = f"/tmp/{file.filename}"
    with open(tmp_path, "wb") as f:
        f.write(await file.read())
    try:
        chunks = preprocess_file(tmp_path, chunk_size)
        for i, chunk_info in enumerate(chunks):
            doc_body = {
                "description": f"{file.filename} - chunk {i+1}",  # title → description
                "content": chunk_info["content"],
                "embedding": chunk_info["embedding"]
            }
            es.index(index=index_name, document=doc_body)
        return {"message": f"{len(chunks)} chunks from file '{file.filename}' indexed successfully."}
    finally:
        os.remove(tmp_path)

# --- ハイブリッド検索エンドポイント ---
# @app.get("/search/")
# def search(index_name: str, query: str, top_k: int = 3):
#     try:
#         query_vector = embedding_model.encode(query).tolist()
#         hybrid_query = {
#             "size": top_k,
#             "query": {
#                 "script_score": {
#                     "query": {
#                         "multi_match": {
#                             "query": query,
#                             "fields": ["description", "content"]
#                         }
#                     },
#                     "script": {
#                         # cosineSimilarity のスコアを multi_match に加点
#                         "source": "cosineSimilarity(params.query_vector, doc['embedding']) + 1.0",
#                         "params": {"query_vector": query_vector}
#                     }
#                 }
#             }
#         }

#         res = es.search(index=index_name, body=hybrid_query)
#         results = [
#             {
#                 "description": hit["_source"].get("description", ""),
#                 "content": hit["_source"].get("content", ""),
#                 "score": hit["_score"]
#             }
#             for hit in res["hits"]["hits"]
#         ]
#         return {"results": results}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
@app.get("/search/")
def search(index_name: str, query: str, top_k: int = 3):
    try:
        query_vector = embedding_model.encode(query).tolist()
        knn_query = {
            "knn": {
                "field": "embedding",
                "query_vector": query_vector,
                "k": top_k,
                "num_candidates": 100
            }
        }

        res = es.search(index=index_name, body=knn_query)
        results = [
            {
                "description": hit["_source"].get("description", ""),
                "content": hit["_source"].get("content", ""),
                "score": hit["_score"]
            }
            for hit in res["hits"]["hits"]
        ]
        return {"results": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- インデックス一覧 ---
@app.get("/list_indices/")
def list_indices():
    try:
        # 全インデックス名取得
        indices_info = es.indices.get(index="*")
        indices_list = []

        for index_name in indices_info.keys():
            # 各インデックスの _meta_ ドキュメント取得
            try:
                meta_res = es.get(index=index_name, id="_meta_")
                description = meta_res["_source"].get("description", "")
            except:
                description = ""  # _meta_ がない場合は空文字

            indices_list.append({"index": index_name, "description": description})

        return {"indices": indices_list}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- インデックス内容取得 ---
@app.get("/index_content/")
def index_content(index_name: str = Query(..., description="取得したいIndex名")):
    try:
        res = es.search(index=index_name, query={"match_all": {}}, size=100)
        documents = [
            {
                "description": hit["_source"].get("description"),
                "content": hit["_source"].get("content"),
                "score": hit["_score"]
            }
            for hit in res["hits"]["hits"]
        ]
        return {"index": index_name, "documents": documents}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- インデックス削除エンドポイント ---
@app.delete("/delete_index/")
def delete_index(index_name: str = Query(..., description="削除したいインデックス名")):
    """
    指定した Elasticsearch インデックスを削除します。
    存在しない場合はエラーを返します。
    """
    try:
        if not es.indices.exists(index=index_name):
            raise HTTPException(status_code=404, detail=f"Index '{index_name}' does not exist.")
        
        res = es.indices.delete(index=index_name)
        if res.get("acknowledged", False):
            return {"message": f"Index '{index_name}' deleted successfully."}
        else:
            raise HTTPException(status_code=500, detail=f"Failed to delete index '{index_name}'.")

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    # uvicorn.run を使って直接サーバー起動
    uvicorn.run(
        "test:app",      # "ファイル名:FastAPIオブジェクト名"
        host="0.0.0.0",
        port=PORT,
        reload=True      # 開発用。コード変更時に自動リロード
    )